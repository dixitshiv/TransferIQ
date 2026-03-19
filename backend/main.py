import os
import uuid
import re
import json as _json
import asyncio
import threading
from io import BytesIO
from datetime import datetime, timezone
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from sse_starlette.sse import EventSourceResponse
from docx import Document as DocxDocument
from docx.shared import Pt

from models.schemas import (
    Transfer, TransferStatus, CreateTransferRequest, DraftRequest,
    Gap, Task, DraftDocument
)
from agents.ingestion import (
    load_demo_package, build_package_summary, generate_demo_pdf,
    generate_package_pdf, DEMO_PDF_PATHS, DEMO_PACKAGE_DIRS, PDF_PATH
)
from utils.pdf_utils import extract_text_from_pdf_bytes
from agents.gap_analysis import run_gap_analysis
from agents.planner import run_planner
from agents.drafter import run_drafter, run_drafter_stream
from agents.rag import retrieve_similar, index_transfer, init_rag
from database import (
    init_db,
    db_get_all_transfers, db_get_transfer, db_transfer_exists, db_create_transfer,
    db_update_transfer_status, db_set_has_demo_data,
    db_get_package, db_has_package, db_set_package_key,
    db_get_gaps, db_has_gaps, db_set_gaps,
    db_get_plan, db_set_plan,
    db_get_draft, db_has_draft, db_set_draft, db_update_draft_content, db_approve_draft,
    db_log_event, db_get_audit_log,
)

app = FastAPI(title="TransferIQ API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


def log_event(tid: str, event_type: str, detail: str):
    db_log_event(tid, event_type, detail, datetime.now(timezone.utc).isoformat())


# ---------------------------------------------------------------------------
# Startup: init DB, seed demo transfers, pre-warm RAG
# ---------------------------------------------------------------------------

DEMO_TRANSFERS = [
    {
        "id": "demo-metformin-001",
        "name": "Metformin HCl 500mg — Site Transfer",
        "product": "Metformin HCl Tablets 500mg",
        "sending_org": "InnoPharm Inc.",
        "receiving_org": "BioMed CDMO",
        "status": TransferStatus.in_progress,
        "has_demo_data": False,
    },
    {
        "id": "demo-ibuprofen-001",
        "name": "Ibuprofen 400mg — Commercial Scale-Up Transfer",
        "product": "Ibuprofen Film-Coated Tablets 400mg",
        "sending_org": "Catalent Pharma Solutions",
        "receiving_org": "Lonza Group AG",
        "status": TransferStatus.in_progress,
        "has_demo_data": False,
    },
    {
        "id": "demo-sitagliptin-001",
        "name": "Sitagliptin Phosphate 100mg — Site Addition",
        "product": "Sitagliptin Phosphate Monohydrate Tablets 100mg",
        "sending_org": "Cambrex Corporation",
        "receiving_org": "Recipharm AB",
        "status": TransferStatus.in_progress,
        "has_demo_data": False,
    },
    {
        "id": "demo-valsartan-001",
        "name": "Valsartan 80mg Capsules — ANDA Commercial Transfer",
        "product": "Valsartan Hard Gelatin Capsules 80mg",
        "sending_org": "Fareva SA",
        "receiving_org": "Thermo Fisher Scientific (Pharma Services)",
        "status": TransferStatus.in_progress,
        "has_demo_data": False,
    },
]


def seed_demo_transfers():
    """Seed demo transfers and their package data into the DB (idempotent)."""
    for t in DEMO_TRANSFERS:
        if db_transfer_exists(t["id"]):
            continue
        db_create_transfer(t)
        # Pre-load the demo package data
        try:
            pkg = load_demo_package(t["id"])
            for key, value in pkg.items():
                db_set_package_key(t["id"], key, value)
            db_set_has_demo_data(t["id"], True)
            log_event(t["id"], "transfer_created", f"Demo transfer '{t['name']}' seeded with package data")
        except Exception as exc:
            log_event(t["id"], "transfer_created", f"Demo transfer '{t['name']}' created (package load error: {exc})")


def generate_demo_pdfs():
    """Generate any missing demo PDFs in the background."""
    for demo_id, pdf_path in DEMO_PDF_PATHS.items():
        if not os.path.exists(pdf_path):
            try:
                generate_package_pdf(demo_id)
            except Exception as exc:
                print(f"[TransferIQ] PDF generation failed for {demo_id}: {exc}")


init_db()
seed_demo_transfers()

# Generate PDFs and pre-warm RAG in background threads
threading.Thread(target=generate_demo_pdfs, daemon=True).start()
threading.Thread(target=init_rag, daemon=True).start()


# ---------------------------------------------------------------------------
# Demo package download
# ---------------------------------------------------------------------------

@app.get("/api/demo-package/download")
def download_demo_package():
    if not os.path.exists(PDF_PATH):
        generate_demo_pdf()
    return FileResponse(
        PDF_PATH,
        media_type="application/pdf",
        filename="demo_transfer_package.pdf",
        headers={"Content-Disposition": "attachment; filename=\"demo_transfer_package.pdf\""},
    )


@app.get("/api/demo-package/{demo_id}/download")
def download_named_demo_package(demo_id: str):
    if demo_id not in DEMO_PDF_PATHS:
        raise HTTPException(404, "Unknown demo package ID")
    pdf_path = DEMO_PDF_PATHS[demo_id]
    if not os.path.exists(pdf_path):
        try:
            generate_package_pdf(demo_id)
        except Exception as exc:
            raise HTTPException(500, f"PDF generation failed: {exc}")
    filename = os.path.basename(pdf_path)
    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Transfers
# ---------------------------------------------------------------------------

@app.get("/api/transfers")
def list_transfers():
    return db_get_all_transfers()


@app.post("/api/transfers", status_code=201)
def create_transfer(req: CreateTransferRequest):
    tid = str(uuid.uuid4())
    t = {
        "id": tid,
        "name": req.name,
        "product": req.product,
        "sending_org": req.sending_org,
        "receiving_org": req.receiving_org,
        "status": TransferStatus.in_progress,
        "has_demo_data": False,
    }
    db_create_transfer(t)
    log_event(tid, "transfer_created", f"Transfer '{req.name}' created")
    return db_get_transfer(tid)


@app.get("/api/transfers/{tid}")
def get_transfer(tid: str):
    t = db_get_transfer(tid)
    if not t:
        raise HTTPException(404, "Transfer not found")
    return t


# ---------------------------------------------------------------------------
# Document upload
# ---------------------------------------------------------------------------

@app.post("/api/transfers/{tid}/upload")
async def upload_docs(tid: str, file: UploadFile = File(...)):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    content = await file.read()
    text, page_count = extract_text_from_pdf_bytes(content)
    db_set_package_key(tid, file.filename, text)
    db_set_has_demo_data(tid, True)
    log_event(tid, "file_uploaded", f"Uploaded '{file.filename}' ({page_count} pages)")
    # Index into RAG vector store
    threading.Thread(target=index_transfer, args=(tid, text[:4000]), daemon=True).start()
    return {"message": f"Uploaded {file.filename}", "pages": page_count}


# ---------------------------------------------------------------------------
# Gap analysis
# ---------------------------------------------------------------------------

@app.post("/api/transfers/{tid}/gap-analysis")
async def run_gap_analysis_stream(tid: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    if not db_has_package(tid):
        raise HTTPException(400, "No documents loaded. Upload a PDF first.")

    package = db_get_package(tid)
    package_summary = build_package_summary(package)
    log_event(tid, "gap_analysis_started", "Gap analysis initiated")

    async def event_generator():
        yield {"data": '{"status": "running", "message": "Analyzing package against ICH requirements..."}'}
        await asyncio.sleep(0)
        try:
            gaps = await asyncio.get_event_loop().run_in_executor(
                None, run_gap_analysis, package_summary
            )
            db_set_gaps(tid, gaps)
            db_update_transfer_status(tid, TransferStatus.gaps_identified)
            log_event(tid, "gap_analysis_complete", f"Gap analysis complete — {len(gaps)} gaps identified")
            import json
            yield {"data": json.dumps({"status": "complete", "gaps": gaps})}
        except Exception as e:
            log_event(tid, "gap_analysis_error", f"Gap analysis failed: {str(e)}")
            yield {"data": f'{{"status": "error", "message": "{str(e)}"}}'}

    return EventSourceResponse(event_generator())


@app.get("/api/transfers/{tid}/gap-analysis")
def get_gap_analysis(tid: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    return db_get_gaps(tid)


# ---------------------------------------------------------------------------
# Transfer plan
# ---------------------------------------------------------------------------

@app.post("/api/transfers/{tid}/plan")
async def generate_plan(tid: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    if not db_has_gaps(tid):
        raise HTTPException(400, "Run gap analysis first.")

    t = db_get_transfer(tid)
    gaps = db_get_gaps(tid)

    tasks = await asyncio.get_event_loop().run_in_executor(
        None, run_planner, t["product"], t["sending_org"], t["receiving_org"], gaps
    )
    db_set_plan(tid, tasks)
    db_update_transfer_status(tid, TransferStatus.plan_ready)
    log_event(tid, "plan_generated", f"Transfer plan generated — {len(tasks)} tasks")
    return tasks


@app.get("/api/transfers/{tid}/plan")
def get_plan(tid: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    return db_get_plan(tid)


# ---------------------------------------------------------------------------
# Draft generation
# ---------------------------------------------------------------------------

@app.post("/api/transfers/{tid}/draft")
async def generate_draft(tid: str, req: DraftRequest):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    if not db_has_package(tid):
        raise HTTPException(400, "No documents loaded.")

    t = db_get_transfer(tid)
    gaps = db_get_gaps(tid)
    package = db_get_package(tid)

    product_info = package.get("product_info") or {
        "product_name": t["product"],
        "sending_org": t["sending_org"],
        "receiving_org": t["receiving_org"],
    }
    log_event(tid, "draft_started", f"Draft generation started: {req.doc_type}")

    async def event_generator():
        import json
        full_content = ""
        async for chunk in run_drafter_stream(req.doc_type, product_info, gaps):
            full_content += chunk
            yield {"data": json.dumps({"chunk": chunk})}
        db_set_draft(tid, req.doc_type, full_content)
        db_update_transfer_status(tid, TransferStatus.draft_complete)
        log_event(tid, "draft_complete", f"Draft complete: {req.doc_type}")
        yield {"data": json.dumps({"status": "complete"})}

    return EventSourceResponse(event_generator())


@app.get("/api/transfers/{tid}/draft/{doc_type}")
def get_draft(tid: str, doc_type: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    draft = db_get_draft(tid, doc_type)
    if not draft:
        raise HTTPException(404, "Draft not found")
    return {"doc_type": doc_type, "content": draft["content"], "status": draft["approval_status"]}


@app.post("/api/transfers/{tid}/draft/{doc_type}/approve")
def approve_draft(tid: str, doc_type: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    if not db_has_draft(tid, doc_type):
        raise HTTPException(404, "Draft not found — generate it first")
    db_approve_draft(tid, doc_type)
    log_event(tid, "draft_approved", f"Draft approved: {doc_type}")
    return {"status": "Approved"}


@app.put("/api/transfers/{tid}/draft/{doc_type}")
async def update_draft(tid: str, doc_type: str, body: dict):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    if not db_has_draft(tid, doc_type):
        raise HTTPException(404, "Draft not found — generate it first")
    db_update_draft_content(tid, doc_type, body.get("content", ""))
    log_event(tid, "draft_edited", f"Draft manually edited: {doc_type}")
    return {"status": "updated"}


# ---------------------------------------------------------------------------
# Draft export (.docx)
# ---------------------------------------------------------------------------

@app.get("/api/transfers/{tid}/draft/{doc_type}/export/docx")
def export_draft_docx(tid: str, doc_type: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    draft = db_get_draft(tid, doc_type)
    if not draft:
        raise HTTPException(404, "Draft not found — generate it first")

    content = draft["content"]
    doc = DocxDocument()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in content.split("\n"):
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("  - ") or stripped.startswith("  * "):
            doc.add_paragraph(stripped[4:], style="List Bullet 2")
        elif stripped == "" or stripped == "---":
            doc.add_paragraph("")
        else:
            para = doc.add_paragraph()
            parts = re.split(r"\*\*(.+?)\*\*", stripped)
            for j, part in enumerate(parts):
                run = para.add_run(part)
                if j % 2 == 1:
                    run.bold = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{doc_type}.docx"
    log_event(tid, "draft_exported", f"Draft exported as .docx: {doc_type}")
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Draft confidence
# ---------------------------------------------------------------------------

@app.get("/api/transfers/{tid}/draft/{doc_type}/confidence")
def get_draft_confidence(tid: str, doc_type: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    draft = db_get_draft(tid, doc_type)
    if not draft:
        raise HTTPException(404, "Draft not found — generate it first")

    content = draft["content"]
    gaps = db_get_gaps(tid)

    sections = re.findall(r'^#{1,3}\s+(.+)$', content, re.MULTILINE)
    if not sections:
        return []

    critical_keywords: set[str] = set()
    major_keywords: set[str] = set()
    for g in gaps:
        words = [w.lower() for w in g.get("category", "").split() if len(w) > 3]
        if g.get("severity") == "CRITICAL":
            critical_keywords.update(words)
        elif g.get("severity") == "MAJOR":
            major_keywords.update(words)

    results = []
    for section in sections:
        section_lower = section.lower()
        if any(kw in section_lower for kw in critical_keywords):
            confidence = "LOW"
            reason = "CRITICAL gap identified in this area — mandatory QA review required before approval"
        elif any(kw in section_lower for kw in major_keywords):
            confidence = "MEDIUM"
            reason = "MAJOR gap identified in this area — reviewer verification recommended"
        else:
            confidence = "HIGH"
            reason = "Adequate source data available in transfer package"
        results.append({"section": section, "confidence": confidence, "reason": reason})

    return results


# ---------------------------------------------------------------------------
# Similar transfers (RAG)
# ---------------------------------------------------------------------------

@app.get("/api/transfers/{tid}/similar")
def get_similar_transfers(tid: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    t = db_get_transfer(tid)
    query = f"{t['product']} transfer from {t['sending_org']} to {t['receiving_org']}"
    package = db_get_package(tid)
    if package:
        query += " " + build_package_summary(package)[:1500]
    return retrieve_similar(query, top_k=5)


# ---------------------------------------------------------------------------
# Audit log
# ---------------------------------------------------------------------------

@app.get("/api/transfers/{tid}/audit-log")
def get_audit_log(tid: str):
    if not db_transfer_exists(tid):
        raise HTTPException(404, "Transfer not found")
    return db_get_audit_log(tid)
